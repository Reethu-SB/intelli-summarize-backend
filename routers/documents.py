"""Router for document upload, text extraction, summarization, and storage.

Endpoints:
  - POST /upload/ : accepts PDF, DOCX, TXT files; saves file to `uploads/`,
    extracts text, summarizes using `services.summarizer`, stores record in DB,
    and returns JSON with `id`, `filename`, and `summary`.

The implementation attempts to use `pdfplumber` for PDFs and falls back to
`PyMuPDF` (`fitz`) if needed. DOCX extraction uses `python-docx`.

The router is defensive: it handles unsupported files, extraction failures,
empty content, and includes an optional retry loop for summary generation.
"""

import os
from io import BytesIO
from typing import Optional
from uuid import uuid4

from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException, status
from sqlalchemy.orm import Session

from database import get_db
from models.document import Document

# Import summarizer service (must expose `summarize_text(text: str, length: str) -> str`)
from services.summarizer import summarize_text


router = APIRouter()

# Directory to store uploaded files
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


def extract_text_from_pdf(data: bytes) -> str:
    """Extract text from PDF bytes. Try pdfplumber first, fallback to PyMuPDF.

    Returns extracted text or raises RuntimeError on failure.
    """
    # Try pdfplumber
    try:
        import pdfplumber

        with pdfplumber.open(BytesIO(data)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n".join(pages).strip()
    except Exception:
        pass

    # Fallback to PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=data, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as exc:
        raise RuntimeError(f"PDF extraction failed: {exc}")


def extract_text_from_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx. Raises RuntimeError on failure."""
    try:
        import docx

        doc = docx.Document(BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise RuntimeError(f"DOCX extraction failed: {exc}")


@router.post("/upload/", status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    length: str = Form("medium"),
    retries: int = Form(1),
    db: Session = Depends(get_db),
):
    """Receive a file, extract text, summarize, store in DB, and return result.

    - `length` chooses summary length: `short`, `medium`, or `long`.
    - `retries` controls how many additional attempts to make if summarization fails.
    """
    # Validate length option
    length = (length or "medium").lower()
    if length not in {"short", "medium", "long"}:
        raise HTTPException(status_code=400, detail="Invalid length; use short|medium|long")

    # Read file bytes
    try:
        contents = await file.read()
    finally:
        await file.close()

    if not contents:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    # Save original file to uploads/ with a unique name
    unique_name = f"{uuid4().hex}_{os.path.basename(file.filename)}"
    save_path = os.path.join(UPLOAD_DIR, unique_name)
    try:
        with open(save_path, "wb") as f:
            f.write(contents)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to save uploaded file: {exc}")

    # Determine file type and extract text
    lower_name = file.filename.lower()
    try:
        if lower_name.endswith(".pdf"):
            text = extract_text_from_pdf(contents)
        elif lower_name.endswith(".docx"):
            text = extract_text_from_docx(contents)
        elif lower_name.endswith(".txt"):
            # Assume utf-8; fall back to latin-1 if utf-8 fails
            try:
                text = contents.decode("utf-8")
            except UnicodeDecodeError:
                text = contents.decode("latin-1")
            text = text.strip()
        else:
            # Unsupported file type
            raise HTTPException(status_code=415, detail="Unsupported file type")
    except RuntimeError as exc:
        # Extraction failure
        raise HTTPException(status_code=422, detail=str(exc))

    if not text:
        raise HTTPException(status_code=422, detail="No extractable text found in the uploaded file")

    # Attempt summarization with retries (retries is number of extra attempts)
    attempt = 0
    summary: Optional[str] = None
    last_error: Optional[Exception] = None
    total_attempts = 1 + max(0, int(retries))
    while attempt < total_attempts:
        try:
            summary = summarize_text(text, length=length)
            break
        except Exception as exc:
            last_error = exc
            attempt += 1

    # If summarization failed after retries, store doc with null summary and return informative response
    if summary is None:
        # Persist original text with empty summary so it can be retried later
        doc = Document(filename=file.filename, content=text, summary=None)
        db.add(doc)
        db.commit()
        db.refresh(doc)
        return {
            "id": doc.id,
            "filename": doc.filename,
            "summary": None,
            "message": f"Summarization failed after {total_attempts} attempts: {last_error}",
        }

    # Save the document and summary to DB
    doc = Document(filename=file.filename, content=text, summary=summary)
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {"id": doc.id, "filename": doc.filename, "summary": summary}
